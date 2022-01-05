import docx
file=open("guests.txt")
names=file.readlines()
doc=docx.Document()
for name in names:
    name=name.strip()
    doc.add_heading("\t\tIt would be a pleasure to have the company of",1)
    doc.add_heading("\t\t\t\t\t"+str(name),1)
    doc.add_heading("\t\t      at 11010 Memory Lane on the Evening of",1)                               
    doc.add_heading("\t\t\t\t\t  April 1st",1)
    doc.add_heading("\t\t\t\t\tat 7 o'clock",1)
    doc.add_page_break()
doc.save("custominvite.docx")
